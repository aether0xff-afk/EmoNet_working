from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
TEX = ROOT / "emonet_paper_v2_ko.tex"
OUT = ROOT / "build_paper_v2_ko" / "emonet_paper_v2_ko_arxiv.docx"
FIG_DIR = ROOT / "figures_full_ko"


BODY_FONT = "Noto Serif KR"
SANS_FONT = "Noto Sans KR"
A4_WIDTH = Inches(8.27)
A4_HEIGHT = Inches(11.69)
TEX_MARGIN = Inches(23 / 25.4)
TEX_TEXT_WIDTH = Inches(8.27 - (2 * 23 / 25.4))


def set_run_font(run, size=None, bold=None, color=None, font_name: str = BODY_FONT):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_style_font(style, size=11, bold=False, color=(0, 0, 0), font_name: str = BODY_FONT):
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(*color)


def clean_inline(text: str) -> str:
    text = text.strip()
    text = text.replace("``", '"').replace("''", '"')
    text = text.replace("\\%", "%").replace("\\_", "_").replace("\\&", "&")
    text = text.replace("--", "-")
    text = re.sub(r"\\eng\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\texttt\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\textit\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\textbf\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\ref\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\label\{[^{}]+\}", "", text)
    text = text.replace("~", " ")
    text = text.replace("\\rightarrow", "→")
    return text.strip()


def add_paragraph(doc: Document, text: str, style=None, align=None):
    if not text.strip():
        return None
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(clean_inline(text))
    set_run_font(run)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.22
    return p


def configure_tex_section(section):
    section.page_width = A4_WIDTH
    section.page_height = A4_HEIGHT
    section.top_margin = TEX_MARGIN
    section.bottom_margin = TEX_MARGIN
    section.left_margin = TEX_MARGIN
    section.right_margin = TEX_MARGIN
    set_columns(section, 1)
    add_footer(section)


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(clean_inline(text), level=level)
    for run in p.runs:
        set_run_font(run, bold=True)
    return p


def set_cell_text(cell, text: str, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(clean_inline(text))
    set_run_font(r, size=9, bold=bold)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=1, cols=width)
    table.style = "Table Grid"
    header = rows[0]
    for i, val in enumerate(header):
        set_cell_text(table.rows[0].cells[i], val, bold=True)
        shade_cell(table.rows[0].cells[i], "E9ECEF")
    for row in rows[1:]:
        cells = table.add_row().cells
        for i in range(width):
            set_cell_text(cells[i], row[i] if i < len(row) else "")
    doc.add_paragraph()


def parse_longtable(lines: list[str], i: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("\\end{longtable}"):
            return rows, i
        if (
            line
            and not line.startswith("\\")
            and "&" in line
            and "\\\\" in line
        ):
            line = line.replace("\\\\", "")
            parts = [clean_inline(p) for p in line.split("&")]
            rows.append(parts)
        i += 1
    return rows, i


def collect_block(lines: list[str], i: int, end_marker: str) -> tuple[list[str], int]:
    block = []
    while i < len(lines):
        if lines[i].strip().startswith(end_marker):
            return block, i
        block.append(lines[i])
        i += 1
    return block, i


def add_figure(doc: Document, block: list[str]):
    fig_name = None
    caption = None
    for line in block:
        m = re.search(r"\\includegraphics(?:\[[^\]]+\])?\{\\figdir/([^{}]+)\}", line)
        if m:
            fig_name = m.group(1)
        c = re.search(r"\\caption\{(.+)\}", line)
        if c:
            caption = clean_inline(c.group(1))
    if not fig_name:
        return
    path = FIG_DIR / fig_name
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=TEX_TEXT_WIDTH)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        set_run_font(cr, size=9.5, color=(55, 55, 55))


def add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("EmoNet: the Emotion Network")
    set_run_font(r, size=9, color=(100, 100, 100))


def set_columns(section, num: int = 1, space_twips: int = 720):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))


def build_docx():
    text = TEX.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    sec = doc.sections[0]
    configure_tex_section(sec)

    set_style_font(doc.styles["Normal"], 11)
    set_style_font(doc.styles["Title"], 22, True, (0, 0, 0))
    set_style_font(doc.styles["Subtitle"], 11, False, (70, 70, 70))
    set_style_font(doc.styles["Heading 1"], 15, True, (0, 0, 0))
    set_style_font(doc.styles["Heading 2"], 12.5, True, (20, 20, 20))
    set_style_font(doc.styles["Heading 3"], 11, True, (40, 40, 40))

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("EmoNet: the Emotion Network")
    set_run_font(r, size=22, bold=True, color=(0, 0, 0))
    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("신경 활성 추적 기반 감정 상태 표현과 에피소드 조건화에 대한 탐구")
    set_run_font(sr, size=11, color=(70, 70, 70))
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("EmoNet Working Draft v2 | 2026년 5월 4일")
    set_run_font(mr, size=9, color=(90, 90, 90))

    i = 0
    in_doc = False
    in_abstract = False
    in_enum = False
    in_quote = False
    in_bib = False
    para_buf: list[str] = []

    def flush():
        nonlocal para_buf
        if para_buf:
            add_paragraph(doc, " ".join(para_buf))
            para_buf = []

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if line.startswith("\\begin{document}"):
            in_doc = True
            i += 1
            continue
        if not in_doc:
            i += 1
            continue
        if line in {"\\onehalfspacing", "\\maketitle"}:
            i += 1
            continue
        if line == "\\tableofcontents":
            flush()
            add_heading(doc, "목차", 1)
            for toc_line in lines:
                section_match = re.match(r"\\section\{(.+)\}", toc_line.strip())
                subsection_match = re.match(r"\\subsection\{(.+)\}", toc_line.strip())
                if section_match:
                    p = add_paragraph(doc, clean_inline(section_match.group(1)))
                    if p:
                        for run in p.runs:
                            run.bold = True
                elif subsection_match:
                    p = add_paragraph(doc, clean_inline("  " + subsection_match.group(1)))
                    if p:
                        p.paragraph_format.left_indent = Inches(0.25)
                        for run in p.runs:
                            set_run_font(run, size=10)
            i += 1
            continue
        if line == "\\newpage":
            flush()
            doc.add_page_break()
            i += 1
            continue
        if line.startswith("\\end{document}"):
            flush()
            break
        if line.startswith("\\begin{abstract}"):
            flush()
            add_heading(doc, "초록", 1)
            in_abstract = True
            i += 1
            continue
        if line.startswith("\\end{abstract}"):
            flush()
            in_abstract = False
            doc.add_section(WD_SECTION.CONTINUOUS)
            body_sec = doc.sections[-1]
            body_sec.footer.is_linked_to_previous = False
            configure_tex_section(body_sec)
            i += 1
            continue
        if line.startswith("\\begin{figure}"):
            flush()
            block, i = collect_block(lines, i + 1, "\\end{figure}")
            add_figure(doc, block)
            i += 1
            continue
        if line.startswith("\\begin{longtable}"):
            flush()
            rows, i = parse_longtable(lines, i + 1)
            add_table(doc, rows)
            i += 1
            continue
        if line.startswith("\\begin{enumerate}"):
            flush()
            in_enum = True
            i += 1
            continue
        if line.startswith("\\end{enumerate}"):
            flush()
            in_enum = False
            i += 1
            continue
        if line.startswith("\\begin{quote}"):
            flush()
            in_quote = True
            i += 1
            continue
        if line.startswith("\\end{quote}"):
            flush()
            in_quote = False
            i += 1
            continue
        if line.startswith("\\begin{thebibliography}"):
            flush()
            in_bib = True
            i += 1
            continue
        if line.startswith("\\end{thebibliography}"):
            flush()
            in_bib = False
            i += 1
            continue
        if line.startswith("\\appendix"):
            flush()
            add_heading(doc, "부록", 1)
            i += 1
            continue
        m = re.match(r"\\section\{(.+)\}", line)
        if m:
            flush()
            add_heading(doc, m.group(1), 1)
            i += 1
            continue
        m = re.match(r"\\subsection\{(.+)\}", line)
        if m:
            flush()
            add_heading(doc, m.group(1), 2)
            i += 1
            continue
        m = re.match(r"\\subsubsection\{(.+)\}", line)
        if m:
            flush()
            add_heading(doc, m.group(1), 3)
            i += 1
            continue
        if line.startswith("\\item"):
            flush()
            item = clean_inline(line.replace("\\item", "", 1))
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(item)
            set_run_font(r)
            i += 1
            continue
        if in_bib and line.startswith("\\bibitem"):
            flush()
            item = re.sub(r"\\bibitem\{[^{}]+\}", "", line).strip()
            if item:
                add_paragraph(doc, item, style="List Bullet")
            i += 1
            continue
        if line.startswith("%") or line.startswith("\\label") or line.startswith("\\toprule") or line.startswith("\\midrule") or line.startswith("\\bottomrule"):
            i += 1
            continue
        if not line:
            flush()
            i += 1
            continue
        if line.startswith("\\"):
            i += 1
            continue
        if in_quote:
            p = add_paragraph(doc, line)
            if p:
                p.paragraph_format.left_indent = Inches(0.3)
                for run in p.runs:
                    run.italic = True
        else:
            para_buf.append(line)
        i += 1

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_docx())
